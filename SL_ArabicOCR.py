from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, PageBreak, Table, TableStyle
import numpy as np
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
from docx.shared import Inches
from docx import Document
from reportlab.lib.units import inch, mm, cm, pica
import logging
from io import BytesIO
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.enums import TA_RIGHT
import re
import os
import pypdfium2 as pdfium
from pdf2image import convert_from_path
from bidi.algorithm import get_display
import arabic_reshaper
from reportlab.pdfbase.ttfonts import TTFont
import reportlab
from reportlab.pdfbase import pdfmetrics
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
import pytesseract
from pdf2image import convert_from_path
import streamlit as st
import base64
from base64 import b64encode
import PyPDF2
from PIL import Image, ImageOps

print("hello to streamlit")
# alias python=/usr/local/bin/python3
pdfmetrics.registerFont(TTFont('Arabic_naskh', 'NotoNaskhArabic-Regular.ttf'))


def show_pdf(file_path):
    with open(file_path, "rb") as f:
        base64_pdf = base64.b64encode(f.read()).decode('utf-8')
    pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="800" height="800" type="application/pdf"></iframe>'
    st.markdown(pdf_display, unsafe_allow_html=True)


if __name__ == '__main__':
    st.markdown("# Arabic Image-to-Text Converter")
    st.markdown(
        " #### This application is used to convert a PDF or image that has Arabic text (scanned or otherwise) and convert it into text that can be searched, copied, analyzed, etc. The output can be downloaded as a PDF or a Word Doc file. \n #### It is IMPORTANT to input how many pages to convert, otherwise it will default to 1.")
    st.markdown(
        "## **Two Options** \n #### 1) Input a .PDF file (recommended)\n #### 2) Input an Image file")
    st.markdown(
        "## Notes: \n - The output will be better if a PDF file is given \n - Once you upload the file, a \"Running\" sign with show at the top right) \n - Once it has completed running, two buttons will appear to download the respective files")
    # Import a PDF in and convert it to images. Save all of the images in a list
    st.markdown(
        "### How many pages would you like to convert?")
    t = st.text_input("Insert amount of pages")
    c1, c2 = st.columns(2)
    with c1:
        pdf_input = st.file_uploader("Input a .PDF")
    with c2:
        img_input = st.file_uploader("Input an Image")

    p2 = pdf_input
    images = []  # This will save all of the images that were converted
    f_name = ""
    #images = convert_from_path(pdffile)
    if pdf_input is not None or img_input is not None:

        if pdf_input is not None:
            f_name = "OCR_" + pdf_input.name[:-4]
            #pdff = PyPDF2.PdfFileReader(open(pdf_input.read(), "rb"))
            if t != "":
                num_of_pages = int(t)
            else:
                num_of_pages = 1
            print(num_of_pages)
            with pdfium.PdfContext(pdf_input) as pdf:
                for page in range(num_of_pages):
                    image = pdfium.render_page_topil(pdf, page)
                    images.append(image)
        else:
            f_name = "OCR_" + img_input.name[:-4]
            pil_img = Image.open(img_input)
            grey = ImageOps.grayscale(pil_img)
            images.append(grey)

        # Run Tesseract to get a list of strings where each string represents a page from the PDF
        # Lil bit less than four minutes for a 148 page PDF
        # 40 minutes for 840 page PDF!

        pages = []  # A list of strings where each string represents a page from the PDF

        print("about to edit and tesseract images")

        # For each image/page, get the text version
        for image in images:
            t = pytesseract.image_to_string(image, lang='ara')
            pages.append(t)

        # Save each string aka page as a page in a PDF
        # only 4 minutes for 840 page PDF !

        # init the style sheet from reportlab
        styles = getSampleStyleSheet()
        styleN = styles['BodyText']
        styleH_1 = styles['Heading1']
        styleH_3 = styles['Heading3']

        a_s_naskh = ParagraphStyle(
            'naskh', fontName='Arabic_naskh', wordWrap='RTL', alignment=2, fontSize=17)
        style_pNum = ParagraphStyle('page_num', alignment=1, fontSize=16)

        # The list that will contain everything that will be ouputted in sequential order
        story = []
        pdf_page = 1

        # RUN THROUGH EVERY PAGE CODE
        story.append(Paragraph(
            "The Quality is NOT 100%. If there is text that is small, it will not accurately portray it", styleH_1))
        for page in pages:
            p = "PDF page: " + str(pdf_page)

            story.append(Paragraph(p, style_pNum))
            story.append(Spacer(1, 12))
            lines = page.split("\n")
            # print(lines)
            for line in lines:
                line = os.linesep.join([s for s in line.splitlines() if s])

                # Run the two functions to reshape the text so it is formatted correctly for the PDF
                reshaped_text = arabic_reshaper.reshape(line)
                line_text = get_display(reshaped_text)

                # make a paragraph of that line with the specific font + settings
                story.append(Paragraph(line_text, a_s_naskh))
                story.append(Spacer(1, 12))

            # Go to the next page since we want to have it correlate to the PDF
            # story.append(PageBreak())
            pdf_page = pdf_page + 1

        # DONE WITH EVERY PAGE CODE

        doc = SimpleDocTemplate(f_name+".pdf", pagesize=letter)
        doc.build(story)  # Create it with all of the information inside
        print('PDF saved Alhamdulillah')

        # Start creation of the Word Doc based on the content from tesseract
        document = Document()
        doc_style = document.styles['Normal']
        doc_font = doc_style.font
        doc_font.size = Pt(15)
        p_num = 1
        document.add_heading('Arabic OCR Reader Word Output', 0)

        for page in pages:
            filtered = filter(lambda x: not re.match(
                r'^\s*$', x), page.split("\n"))
            ppage = list(filtered)
            pp = document.add_paragraph("PDF Page: " + str(p_num) + "\n")
            for i in range(len(ppage)):
                p_reshape = arabic_reshaper.reshape(ppage[i])
                p = document.add_paragraph(p_reshape)
                p.style.font.rtl = True
                p.style.font.size = Pt(17)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_num += 1
            # document.add_page_break()
        document.save(f_name+".docx")
        print("Word doc saved alhamdulillah")

        with open(f_name+".pdf", "rb") as pdf_file:
            PDFbyte = pdf_file.read()
        with open(f_name+".docx", "rb") as word_file:
            wordByte = word_file.read()
        col1, col2, col3 = st.columns(3)

        with col1:
            pass
        with col2:
            st.download_button(label="Download the .PDF file", data=PDFbyte,
                               file_name="test.pdf", mime='application/octet-stream')
            st.download_button("Download the .Docx file", data=wordByte,
                               file_name=f_name+".docx", mime='application/octet-stream')
        with col3:
            pass
        # show_pdf(f_name+".pdf")
    st.markdown("")
    st.markdown("")
    st.markdown(
        "[Coded up by Syed Raza!](https://ssraza21.github.io/index.html)")
